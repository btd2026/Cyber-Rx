# -*- coding: utf-8 -*-
"""
guide_kit.py — reusable McKinsey-style .docx toolkit for the Nerion manual.
Import this from the content script: `from guide_kit import *; K = Kit()`.
Provides cover, 3-level numbered headings (all TOC-linked), exhibit tables,
field-reference tables, step procedures with expected results, and note/tip/
warning callouts, plus running header/footer with page numbers.
"""
import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY   = RGBColor(0x08, 0x24, 0x3A)
BLUE   = RGBColor(0x1A, 0x5F, 0xA0)
CYAN   = RGBColor(0x2F, 0xA3, 0xD6)
INK    = RGBColor(0x22, 0x2A, 0x30)
MUTE   = RGBColor(0x5B, 0x6B, 0x76)
GOOD   = RGBColor(0x1E, 0x7A, 0x3C)
WARN   = RGBColor(0xB0, 0x6A, 0x00)
CRIT   = RGBColor(0xA6, 0x25, 0x25)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
NAVYHX="08243A"; BLUEHX="1A5FA0"; CYANHX="2FA3D6"
LIGHT="EEF2F5"; BAND="E7EEF3"; TIPHX="ECF5EF"; WARNHX="FBF1E4"; NOTEHX="EEF2F5"
HEAD="Georgia"; BODY="Calibri"; MONO="Consolas"

def _sh(el,hx):
    s=OxmlElement('w:shd'); s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),hx); el.append(s)
def psh(p,hx): _sh(p._p.get_or_add_pPr(),hx)
def csh(c,hx): _sh(c._tc.get_or_add_tcPr(),hx)
def bdr(p,edges,sz=6,color=BLUEHX,space='6'):
    pPr=p._p.get_or_add_pPr(); b=OxmlElement('w:pBdr')
    for e in edges:
        x=OxmlElement('w:'+e); x.set(qn('w:val'),'single'); x.set(qn('w:sz'),str(sz)); x.set(qn('w:space'),space); x.set(qn('w:color'),color); b.append(x)
    pPr.append(b)
def _field(p,instr):
    r=p.add_run(); a=OxmlElement('w:fldChar'); a.set(qn('w:fldCharType'),'begin'); r._r.append(a)
    t=OxmlElement('w:instrText'); t.set(qn('xml:space'),'preserve'); t.text=instr; r._r.append(t)
    s=OxmlElement('w:fldChar'); s.set(qn('w:fldCharType'),'separate'); r._r.append(s)
    e=OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'),'end'); r._r.append(e)

class Kit:
    def __init__(self):
        self.doc=Document(); self.sec=[0]; self.sub=[0]; self.ssub=[0]
        n=self.doc.styles["Normal"]; n.font.name=BODY; n.font.size=Pt(10.5); n.font.color.rgb=INK
        n.paragraph_format.space_after=Pt(6); n.paragraph_format.line_spacing=1.13
        for s in self.doc.sections:
            s.top_margin=Inches(0.9); s.bottom_margin=Inches(0.9); s.left_margin=Inches(1.0); s.right_margin=Inches(1.0)

    def run(self,p,text,size=10.5,bold=False,italic=False,color=INK,font=BODY,caps=False,spacing=None,mono=False):
        r=p.add_run(text); r.font.name=(MONO if mono else font); r.font.size=Pt(size); r.font.bold=bold; r.font.italic=italic; r.font.color.rgb=color
        if caps: r.font.all_caps=True
        if spacing is not None:
            sp=OxmlElement('w:spacing'); sp.set(qn('w:val'),str(spacing)); r._r.get_or_add_rPr().append(sp)
        return r
    def p(self,text=None,before=0,after=6,align=None,indent=None,**kw):
        pp=self.doc.add_paragraph(); pp.paragraph_format.space_before=Pt(before); pp.paragraph_format.space_after=Pt(after)
        if align is not None: pp.alignment=align
        if indent: pp.paragraph_format.left_indent=Inches(indent)
        if text: self.run(pp,text,**kw)
        return pp
    def _outline(self,p,lvl):
        o=OxmlElement('w:outlineLvl'); o.set(qn('w:val'),str(lvl)); p._p.get_or_add_pPr().append(o)
    def h1(self,title):
        self.sec[0]+=1; self.sub[0]=0
        self.doc.add_page_break()
        a=self.doc.add_paragraph(); a.paragraph_format.space_after=Pt(0); self.run(a,f"CHAPTER {self.sec[0]:02d}",size=9,bold=True,color=CYAN,caps=True,spacing=40)
        b=self.doc.add_paragraph(); b.paragraph_format.space_after=Pt(10); self.run(b,title,size=22,bold=True,color=NAVY,font=HEAD); bdr(b,['bottom'],sz=12,color=NAVYHX); self._outline(b,0)
        return b
    def h2(self,title):
        self.sub[0]+=1; self.ssub[0]=0
        p=self.doc.add_paragraph(); p.paragraph_format.space_before=Pt(13); p.paragraph_format.space_after=Pt(3)
        self.run(p,f"{self.sec[0]}.{self.sub[0]}  ",size=14,bold=True,color=CYAN,font=HEAD); self.run(p,title,size=14,bold=True,color=BLUE,font=HEAD); self._outline(p,1)
        return p
    def h3(self,title):
        self.ssub[0]+=1
        p=self.doc.add_paragraph(); p.paragraph_format.space_before=Pt(9); p.paragraph_format.space_after=Pt(2)
        self.run(p,title,size=11.5,bold=True,color=NAVY); self._outline(p,2)
        return p
    def bullet(self,text,lead=None,level=0):
        p=self.doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(2); p.paragraph_format.left_indent=Inches(0.3+0.25*level)
        if lead: self.run(p,lead,bold=True,color=NAVY); self.run(p,text)
        else: self.run(p,text)
        return p
    def steps(self,items):
        # items: list of (title, body) or (title, body, expected)
        for i,it in enumerate(items,1):
            title=it[0]; body=it[1] if len(it)>1 else None; exp=it[2] if len(it)>2 else None
            p=self.doc.add_paragraph(); p.paragraph_format.space_before=Pt(5); p.paragraph_format.space_after=Pt(1)
            self.run(p,f"{i}  ",size=11.5,bold=True,color=CYAN,font=HEAD); self.run(p,title,size=11,bold=True,color=NAVY)
            if body:
                b=self.doc.add_paragraph(); b.paragraph_format.left_indent=Inches(0.30); b.paragraph_format.space_after=Pt(2); self.run(b,body)
            if exp:
                e=self.doc.add_paragraph(); e.paragraph_format.left_indent=Inches(0.30); e.paragraph_format.space_after=Pt(4)
                self.run(e,"Result:  ",size=9.5,bold=True,color=GOOD); self.run(e,exp,size=9.5,italic=True,color=MUTE)
    def callout(self,kind,text,label=None):
        cfg={'note':(NOTEHX,BLUEHX,BLUE,'Note'),'tip':(TIPHX,'1E7A3C',GOOD,'Tip'),
             'warn':(WARNHX,'B06A00',WARN,'Important'),'key':(LIGHT,BLUEHX,BLUE,'Key takeaway')}
        fill,bhx,tc,dl=cfg.get(kind,cfg['note'])
        p=self.doc.add_paragraph(); p.paragraph_format.space_before=Pt(7); p.paragraph_format.space_after=Pt(8); p.paragraph_format.left_indent=Inches(0.06)
        psh(p,fill); bdr(p,['left'],sz=24,color=bhx)
        self.run(p,(label or dl).upper()+"   ",size=8.5,bold=True,color=tc,caps=True,spacing=25); self.run(p,text,size=10,color=INK)
        return p
    def table(self,caption,headers,rows,widths=None,cap_kicker="EXHIBIT"):
        if caption:
            c=self.doc.add_paragraph(); c.paragraph_format.space_before=Pt(10); c.paragraph_format.space_after=Pt(3)
            self.run(c,cap_kicker+" — ",size=8.5,bold=True,color=CYAN,caps=True,spacing=18); self.run(c,caption,size=8.5,bold=True,color=MUTE,caps=True,spacing=18)
        t=self.doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER
        hc=t.rows[0].cells
        for i,h in enumerate(headers):
            csh(hc[i],NAVYHX); cp=hc[i].paragraphs[0]; cp.paragraph_format.space_before=Pt(2); cp.paragraph_format.space_after=Pt(2)
            self.run(cp,h,size=9,bold=True,color=WHITE)
        for ri,r in enumerate(rows):
            cells=t.add_row().cells
            for ci,val in enumerate(r):
                if ri%2==1: csh(cells[ci],BAND)
                cp=cells[ci].paragraphs[0]; cp.paragraph_format.space_before=Pt(1); cp.paragraph_format.space_after=Pt(1)
                lead=(ci==0)
                self.run(cp,str(val),size=9,bold=lead,color=(NAVY if lead else INK))
        if widths:
            for i,w in enumerate(widths):
                for row in t.rows: row.cells[i].width=Inches(w)
        self.doc.add_paragraph().paragraph_format.space_after=Pt(2)
        return t
    def fields(self,caption,rows,widths=(1.7,1.0,3.8)):
        # rows: (field, type, description)
        return self.table(caption,["Field","Type","What it does / options"],rows,widths=list(widths),cap_kicker="FIELDS")
    def save(self,path):
        self.doc.save(path); return path
