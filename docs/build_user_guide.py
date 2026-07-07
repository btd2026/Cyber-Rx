#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Nerion User Guide — McKinsey-style Word document generator.
Produces a polished, editable .docx: cover, auto TOC, numbered sections,
exhibit tables, key-takeaway callouts, running header/footer with page numbers.
"""
import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- Palette (deep navy consulting + Nerion blue) ----------
NAVY   = RGBColor(0x08, 0x24, 0x3A)   # headings, cover
BLUE   = RGBColor(0x1A, 0x5F, 0xA0)   # Nerion accent
CYAN   = RGBColor(0x2F, 0xA3, 0xD6)   # bright accent / rules
INK    = RGBColor(0x22, 0x2A, 0x30)   # body text
MUTE   = RGBColor(0x5B, 0x6B, 0x76)   # captions
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT  = "EEF2F5"                       # shaded fills (hex string)
NAVYHX = "08243A"
BLUEHX = "1A5FA0"
BANDHX = "E7EEF3"

HEAD_FONT = "Georgia"      # serif headings — authoritative report feel
BODY_FONT = "Calibri"      # clean sans body

doc = Document()

# ---------- base styles ----------
normal = doc.styles["Normal"]
normal.font.name = BODY_FONT
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12

# default margins
for s in doc.sections:
    s.top_margin = Inches(0.9); s.bottom_margin = Inches(0.9)
    s.left_margin = Inches(1.0); s.right_margin = Inches(1.0)

# ---------- low-level helpers ----------
def _shade(el, hex_fill):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_fill)
    el.append(shd)

def _p_shade(p, hex_fill):
    _shade(p._p.get_or_add_pPr(), hex_fill)

def _cell_shade(cell, hex_fill):
    _shade(cell._tc.get_or_add_tcPr(), hex_fill)

def _borders(p, edges, sz=6, color=BLUEHX):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    for edge in edges:
        e = OxmlElement('w:'+edge)
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), str(sz))
        e.set(qn('w:space'), '6'); e.set(qn('w:color'), color)
        pbdr.append(e)
    pPr.append(pbdr)

def run(p, text, *, size=10.5, bold=False, italic=False, color=INK, font=BODY_FONT, caps=False, spacing=None):
    r = p.add_run(text)
    r.font.name = font; r.font.size = Pt(size); r.font.bold = bold; r.font.italic = italic
    r.font.color.rgb = color
    if caps: r.font.all_caps = True
    if spacing is not None:
        rPr = r._r.get_or_add_rPr(); sp = OxmlElement('w:spacing'); sp.set(qn('w:val'), str(spacing)); rPr.append(sp)
    return r

def para(text=None, *, before=0, after=6, align=None, **kw):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
    if align is not None: p.alignment = align
    if text: run(p, text, **kw)
    return p

# ---------- section numbering ----------
SEC = [0]
def h1(title):
    SEC[0]+= 1
    doc.add_page_break()
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    run(p, f"{SEC[0]:02d}", size=30, bold=True, color=CYAN, font=HEAD_FONT)
    p2 = doc.add_paragraph(); p2.paragraph_format.space_after = Pt(10)
    run(p2, title, size=21, bold=True, color=NAVY, font=HEAD_FONT)
    _borders(p2, ['bottom'], sz=10, color=NAVYHX)
    p2.style = p2.style  # keep in TOC via outline level
    _outline(p2, 1)
    return p2

SUB = [0]
def h2(title):
    SUB[0]+=1
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    run(p, title, size=14, bold=True, color=BLUE, font=HEAD_FONT)
    _outline(p, 2)
    return p

def h3(title):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
    run(p, title, size=11.5, bold=True, color=NAVY, font=BODY_FONT)
    _outline(p, 3)
    return p

def _outline(p, lvl):
    pPr = p._p.get_or_add_pPr()
    o = OxmlElement('w:outlineLvl'); o.set(qn('w:val'), str(lvl-1)); pPr.append(o)

def bullet(text, *, bold_lead=None, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.3 + 0.25*level)
    if bold_lead:
        run(p, bold_lead, bold=True, color=NAVY); run(p, text)
    else:
        run(p, text)
    return p

def numbered(n, title, body):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    run(p, f"{n}  ", size=12, bold=True, color=CYAN, font=HEAD_FONT)
    run(p, title, size=12, bold=True, color=NAVY, font=BODY_FONT)
    if body:
        b = doc.add_paragraph(); b.paragraph_format.left_indent = Inches(0.32); b.paragraph_format.space_after = Pt(6)
        run(b, body)
    return p

def callout(label, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.08)
    _p_shade(p, LIGHT); _borders(p, ['left'], sz=24, color=BLUEHX)
    run(p, label.upper()+"   ", size=8.5, bold=True, color=BLUE, caps=True, spacing=30)
    run(p, text, size=10, italic=False, color=INK)
    # padding
    p.paragraph_format.space_after = Pt(8)
    return p

def exhibit(caption, headers, rows, widths=None):
    cap = doc.add_paragraph(); cap.paragraph_format.space_before = Pt(10); cap.paragraph_format.space_after = Pt(3)
    run(cap, "EXHIBIT — ", size=8.5, bold=True, color=CYAN, caps=True, spacing=20)
    run(cap, caption, size=8.5, bold=True, color=MUTE, caps=True, spacing=20)
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    # header row
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        _cell_shade(hdr[i], NAVYHX)
        cp = hdr[i].paragraphs[0]; cp.paragraph_format.space_after = Pt(2); cp.paragraph_format.space_before = Pt(2)
        run(cp, htext, size=9.5, bold=True, color=WHITE, font=BODY_FONT)
    # body
    for ri, r in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(r):
            if ri % 2 == 1: _cell_shade(cells[ci], BANDHX)
            cp = cells[ci].paragraphs[0]; cp.paragraph_format.space_after = Pt(2); cp.paragraph_format.space_before = Pt(2)
            lead_bold = (ci == 0)
            run(cp, str(val), size=9.5, bold=lead_bold, color=(NAVY if lead_bold else INK))
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

# ---------- header / footer with page numbers ----------
def add_field(p, instr):
    r = p.add_run()
    fb = OxmlElement('w:fldChar'); fb.set(qn('w:fldCharType'), 'begin'); r._r.append(fb)
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = instr; r._r.append(it)
    fs = OxmlElement('w:fldChar'); fs.set(qn('w:fldCharType'), 'separate'); r._r.append(fs)
    fe = OxmlElement('w:fldChar'); fe.set(qn('w:fldCharType'), 'end'); r._r.append(fe)

# ---------------------------------------------------------------------------
# COVER PAGE
# ---------------------------------------------------------------------------
# top navy rule
top = doc.add_paragraph(); top.paragraph_format.space_after = Pt(0)
_p_shade(top, NAVYHX); run(top, " ", size=2)
for _ in range(1): doc.add_paragraph()

lbl = doc.add_paragraph(); lbl.paragraph_format.space_after = Pt(2)
run(lbl, "NERION", size=13, bold=True, color=NAVY, font=HEAD_FONT, spacing=60)
sub = doc.add_paragraph(); sub.paragraph_format.space_after = Pt(40)
run(sub, "Cyber Business Operations Platform", size=11, color=BLUE, font=BODY_FONT, spacing=20)

t1 = doc.add_paragraph(); t1.paragraph_format.space_after = Pt(2)
run(t1, "User Guide", size=40, bold=True, color=NAVY, font=HEAD_FONT)
t2 = doc.add_paragraph(); t2.paragraph_format.space_after = Pt(24)
run(t2, "From onboarding to daily operation", size=16, italic=True, color=MUTE, font=HEAD_FONT)

rule = doc.add_paragraph(); rule.paragraph_format.space_after = Pt(18)
_borders(rule, ['bottom'], sz=18, color=BLUEHX)

lede = doc.add_paragraph(); lede.paragraph_format.space_after = Pt(6)
run(lede, "Give Nerion your business, and it returns your cyber position — in dollars. "
          "This guide walks every executive seat from first login through onboarding, "
          "the cockpit, decisions, and administration of the platform.",
    size=11.5, color=INK)

for _ in range(8): doc.add_paragraph()

meta = doc.add_paragraph(); meta.paragraph_format.space_after = Pt(1)
run(meta, "Prepared for platform users and administrators", size=9.5, color=MUTE)
meta2 = doc.add_paragraph(); meta2.paragraph_format.space_after = Pt(1)
run(meta2, "Version 1.0  ·  " + datetime.date.today().strftime("%B %Y"), size=9.5, color=MUTE)
conf = doc.add_paragraph()
run(conf, "CONFIDENTIAL — Distribution limited to licensed users", size=8.5, bold=True, color=BLUE, caps=True, spacing=20)

# bottom navy rule
bot = doc.add_paragraph(); bot.paragraph_format.space_before = Pt(6)
_p_shade(bot, NAVYHX); run(bot, " ", size=2)

# ---------------------------------------------------------------------------
# TABLE OF CONTENTS
# ---------------------------------------------------------------------------
doc.add_page_break()
toch = doc.add_paragraph(); toch.paragraph_format.space_after = Pt(8)
run(toch, "Contents", size=20, bold=True, color=NAVY, font=HEAD_FONT)
_borders(toch, ['bottom'], sz=10, color=NAVYHX)
tp = doc.add_paragraph()
add_field(tp, r'TOC \o "1-2" \h \z \u')
hint = doc.add_paragraph();
run(hint, "To populate: right-click anywhere in the contents above and choose "
          "“Update Field” (or select it and press F9).", size=8.5, italic=True, color=MUTE)

# set running footer on the main section (applies document-wide here)
sec = doc.sections[0]
footer = sec.footer
fp = footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(fp, "Nerion User Guide   ·   Confidential      ", size=8, color=MUTE)
add_field(fp, "PAGE")
run(fp, " ", size=8, color=MUTE)
# header
hp = sec.header.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run(hp, "NERION", size=8, bold=True, color=BLUE, spacing=30)

# ===========================================================================
# 1 · INTRODUCTION
# ===========================================================================
h1("Introduction")

h2("What Nerion is")
para("Nerion is a Cyber Business Operations Platform. It takes what your business is — "
     "your revenue processes, the systems that run them, your crown-jewel data, your vendors, "
     "and your security tools — and returns your cyber position expressed the way the board "
     "and the C-suite think: in dollars of exposure, in decisions, and in defensible evidence.")
para("Unlike a dashboard that shows red and green tiles, every number in Nerion is traceable. "
     "Each figure carries its provenance — the inputs, the calculation, and the source signal it "
     "came from — so a metric is never a claim you have to take on faith. Where a figure is drawn "
     "from live data it is marked as live; where it is a modelled placeholder it is marked as "
     "illustrative.")

h2("Who it is for")
para("Nerion presents one platform through ten executive lenses, called seats. Each seat is "
     "addressed to a named leader, speaks in that leader’s language, and stamps that leader’s name "
     "on every decision they make. The seats are summarized in Exhibit 1.")

exhibit("The ten executive seats",
    ["Seat", "Owner", "What the seat answers"],
    [
        ["CEO", "Chief Executive", "Is the business healthy, governed, and disclosable?"],
        ["CISO", "Security", "Where is exposure, and are controls actually working?"],
        ["CFO", "Finance", "What is the dollar exposure, insurance gap, and ROI?"],
        ["CRO", "Chief Risk", "How does cyber sit beside our other principal risks?"],
        ["CTO", "Technology", "Which systems carry the business, and their app-sec health?"],
        ["CLO", "General Counsel", "Materiality, DSARs, litigation holds, obligations."],
        ["COO", "Operations", "Resilience, recovery times, business continuity."],
        ["CPO", "Product", "Secure-by-design coverage and product data risk."],
        ["Internal Audit", "Chief Audit Exec", "Audit plan, workpapers, control assurance."],
        ["Board", "Directors", "Oversight, disclosure readiness, peer posture."],
    ],
    widths=[1.1, 1.5, 3.9])

callout("The core idea",
    "You describe the business once during onboarding. Nerion maps process → system → "
    "crown jewel → risk → control, prices the exposure, and lights up all ten seats from that "
    "single description. The more you provide, the more of the cockpit is live rather than illustrative.")

h2("How to use this guide")
para("Sections 2–3 take a new customer from first login through onboarding. Section 4 orients you "
     "in the cockpit. Section 5 is a reference for each seat. Section 6 covers the day-to-day "
     "workflows every user performs. Sections 7–8 cover administration, licensing, and support.")

# ===========================================================================
# 2 · GETTING STARTED
# ===========================================================================
h1("Getting Started")

h2("How Nerion is delivered")
para("Nerion is delivered as a self-contained virtual appliance that runs inside your own "
     "environment. Your data never leaves your network — there is no phone-home and no external "
     "dependency for the platform to operate. Your administrator receives the appliance image and "
     "a license file; once the appliance is running, users reach Nerion through a browser.")

exhibit("License types",
    ["Plan", "Term", "Intended use"],
    [
        ["Trial", "14 days", "Evaluation with your own data, full functionality."],
        ["Paid subscription", "365 days", "Production use; renews annually."],
    ], widths=[1.6, 1.4, 3.5])

para("The platform shows the remaining term in-app. As the expiry approaches, a banner prompts "
     "renewal; a short grace period after expiry keeps the platform readable while your renewal is "
     "processed. Administration of the license is covered in Section 7.")

h2("First login")
numbered("1", "Open the platform", "Navigate to the address your administrator provides. "
    "Nerion runs in any modern browser; no plug-in is required.")
numbered("2", "Choose your seat", "Select the executive lens you own — for example CISO or CFO. "
    "You can switch seats at any time from the sidebar; the platform is one system viewed ten ways.")
numbered("3", "Start onboarding", "If your organization has not yet been described, you will be "
    "taken into the onboarding wizard (Section 3). If setup is complete, you land directly in the cockpit.")

callout("Everything can be edited later",
    "Onboarding fields are optional and reversible. Provide what you know now; you can refine every "
    "input afterward and the cockpit will recompute. Nothing you enter is locked.")

# ===========================================================================
# 3 · ONBOARDING
# ===========================================================================
h1("Onboarding: Describing Your Business")

para("Onboarding is a guided, single-page wizard. Its purpose is one sentence: give Nerion your "
     "business so it can return your cyber position in dollars. You move through the sections in "
     "Exhibit 2; each feeds specific parts of the cockpit. None is mandatory, but each one you "
     "complete converts more of the cockpit from illustrative to live.")

exhibit("The onboarding sections, in order",
    ["#", "Section", "What it establishes"],
    [
        ["1", "Organization", "Industry lens, currency, regulators and jurisdictions."],
        ["2", "Leadership", "The named owner of each seat and each decision."],
        ["3", "Business processes", "Your revenue-bearing processes — the basis of crown jewels."],
        ["4", "Systems & applications", "The systems that run those processes."],
        ["5", "Third-party vendors", "Tier-1/2 vendors and concentration risk."],
        ["6", "Connect your tools", "Live signal feeds from your security stack."],
        ["7", "Board governance", "Oversight cadence and incident-readiness (SEC Item 106)."],
        ["8", "AI risk & governance", "AI systems you run and AI you use to defend."],
        ["9", "Strategic objectives", "Business goals cyber risk is measured against."],
        ["10", "Other principal risks", "Credit, market, operational — for the CRO one-scale view."],
        ["→", "Go live", "Submit; the cockpit builds from your description."],
    ], widths=[0.4, 1.9, 4.2])

h2("Step 1 — Organization")
para("Set your organization name, industry, reporting currency, and operating regions. Industry "
     "selects the peer lens and the regulatory frame. Operating regions set the notification clocks "
     "and penalty regimes that apply — for example, choosing EU and UK activates GDPR and UK data-"
     "protection timelines used later in the CLO and Board seats.")
callout("Why region matters", "Region drives real deadlines. It is used to compute breach-"
    "notification countdowns and maximum penalties, so set every region in which you operate.")

h2("Step 2 — Leadership")
para("Enter the name of each executive who owns a seat: CEO, CFO, CISO, COO, CTO, CPO, General "
     "Counsel (CLO), Chief Risk Officer, and Chief Audit Executive. These names are not cosmetic — "
     "each seat’s cockpit is addressed to its owner, and every decision recorded in Nerion is "
     "stamped with the name of the executive who made it, forming an accountable record.")

h2("Step 3 — Business processes")
para("List your key revenue-bearing processes (for example, Payments, Claims, Trading, Order "
     "Fulfilment). You can type them or upload from a CMDB export. Nerion groups processes by "
     "business function and derives your crown jewels — the concentrations of value an attacker "
     "would target — from them. This is the single most valuable section: it anchors the "
     "process → system → crown-jewel → risk → control chain that the whole platform is built on.")

h2("Step 4 — Systems & applications")
para("Map the systems and applications that run each process. For each system you can note its "
     "exposure (for example internet-facing) and criticality. Exposure and criticality drive the "
     "attack-surface and crown-jewel calculations, so they are worth setting accurately. Systems "
     "can be entered manually or imported.")

h2("Step 5 — Third-party vendors")
para("Capture your tier-1 and tier-2 vendors, by CSV upload or from a third-party risk API. Nerion "
     "uses this to surface concentration risk (over-reliance on a single provider) and to populate "
     "the Third-party risk views in the CEO, CISO, and CRO seats.")

h2("Step 6 — Connect your tools")
para("Connect your security and operations tools so the cockpit reads live signals instead of "
     "modelled placeholders. Nerion ships native connectors across the stack; a representative set "
     "is shown in Exhibit 3. Each connector maps to specific metrics in the seats.")

exhibit("Representative native connectors",
    ["Category", "Example tools", "Signals fed"],
    [
        ["Cloud posture", "Wiz / CSPM", "Cloud security posture score"],
        ["Email security", "Proofpoint TAP", "BEC / phishing blocked"],
        ["Privacy / DSAR", "OneTrust", "Open & overdue DSARs, legal holds"],
        ["Threat intel", "Recorded Future", "Active threat actors"],
        ["Backup / DR", "Rubrik", "Immutable backups, DR test age, RPO"],
        ["SOX / GRC", "SAP GRC", "SoD conflicts, change control, payment anomalies"],
        ["Identity", "SailPoint / CyberArk", "Dormant accounts, flagged privileged sessions"],
    ], widths=[1.5, 2.0, 3.0])

callout("Live vs illustrative", "Any metric with a connected source is labelled live and updates "
    "on refresh. Metrics without a source remain illustrative so you are never misled about which "
    "numbers are grounded in your data.")

h2("Step 7 — Board governance & incident readiness")
para("Record how the board oversees cyber: the committee that owns it, meeting cadence, and your "
     "incident-response readiness. These inputs power the CEO governance panel and the Board seat’s "
     "SEC Item 106 / 10-K disclosure-readiness view.")

h2("Step 8 — AI risk & governance")
para("Nerion treats AI in two halves: the AI you run (systems that could be attacked or misused) "
     "and AI you use to defend. Register your AI/LLM systems, the governance framework you follow "
     "(for example NIST AI RMF), and whether the EU AI Act is in scope. This powers the AI-risk "
     "sections in the CEO and CISO seats and the continuous AI-framework audit.")

h2("Step 9 — Strategic objectives")
para("Enter the strategic objectives the business is pursuing. Nerion measures cyber risk against "
     "them so the CEO seat can express security in terms of the goals leadership actually cares "
     "about, rather than in isolation.")

h2("Step 10 — Other principal risks")
para("Provide rough magnitudes for your other principal risks — credit/market, operational, "
     "compliance/legal, third-party. This lets the CRO seat place cyber on one comparable scale "
     "beside the enterprise’s other risks, which is how risk committees actually weigh it.")

h2("Go live")
para("Select Go live to submit. Nerion ingests your description, runs the crown-jewel and economics "
     "engines, and builds the cockpit. You are taken to your seat with the platform populated from "
     "your data. Return to onboarding at any time to refine inputs.")

# ===========================================================================
# 4 · THE COCKPIT
# ===========================================================================
h1("The Cockpit: Orientation")

h2("Layout")
para("The cockpit has three constants. The seat sidebar (left) switches between the ten executive "
     "lenses. The seat header (top) shows whose cockpit you are in and an ‘as of’ timestamp for the "
     "data. The tab bar (below the header) organizes that seat’s content. Selecting a different seat "
     "re-addresses the entire cockpit to that leader.")

h2("Reading a metric: provenance")
para("Every headline number is backed by the provenance engine. Selecting a metric opens its "
     "inspection: the inputs that fed it, the calculation applied, and the live signal or onboarding "
     "value it derives from. This is what makes Nerion defensible in front of an auditor or a "
     "regulator — no figure is a black box.")

exhibit("The four things every metric tells you",
    ["Element", "Meaning"],
    [
        ["Value", "The current figure, in your reporting currency where applicable."],
        ["Direction", "Whether it is improving or deteriorating, and over what period."],
        ["Provenance", "The inputs, calculation, and source signal behind the value."],
        ["Live / illustrative", "Whether it is grounded in connected data or modelled."],
    ], widths=[1.8, 4.7])

h2("Decisions")
para("A decision is the unit of action in Nerion. When a seat surfaces a choice — accept a risk, "
     "fund an initiative, escalate a materiality question — you record the decision with its "
     "rationale. It is stamped with your name and the time. Decisions remain editable for a short "
     "window (24 hours) and then commit to the permanent record. Committed decisions can be pushed "
     "to your ticketing system so execution is tracked where your teams already work (Section 6).")

h2("Evidence")
para("Where a seat makes a claim about control effectiveness or exposure treated, the underlying "
     "evidence is one selection away. Evidence links the business claim back to the control, the "
     "signal, and the calculation — closing the loop between what leadership is told and what is "
     "actually true in the environment.")

# ===========================================================================
# 5 · SEAT-BY-SEAT
# ===========================================================================
h1("Seat-by-Seat Reference")

para("Each seat below lists what it answers and the panels you will use most. All ten draw from the "
     "same onboarding description and live signals; they differ in lens, not in underlying data.")

h2("CEO — Chief Executive")
bullet("business health measured against your strategic objectives.", bold_lead="Business health:  ")
bullet("governance cadence and SEC Item 106 / 10-K disclosure readiness.", bold_lead="Governance & oversight:  ")
bullet("concentration and named stress scenarios across tier-1 vendors.", bold_lead="Third-party concentration:  ")
bullet("recovery readiness and a named enterprise stress scenario.", bold_lead="Incident & recovery:  ")
bullet("posture versus published peer medians for your industry.", bold_lead="Peer benchmark:  ")

h2("CISO — Security")
para("The CISO seat is organized into six tabs, the deepest in the platform:")
exhibit("CISO tabs",
    ["Tab", "What it covers"],
    [
        ["Program Health", "Direction of the program, capability coverage, and active-compromise status."],
        ["Threats", "Live threat status and coverage drawn from connected SIEM/EDR signals."],
        ["My Decisions", "The CISO’s recorded decisions, editable within 24h then committed."],
        ["AI Risk", "The AI you run and the AI you use to defend — with continuous framework audit."],
        ["Third-party", "Top vendor exposures with live scores on one scale."],
        ["Frameworks", "Control assessment across NIST CSF, 800-53, SOC 2, HIPAA, CIS — with drill-down."],
    ], widths=[1.6, 4.9])
para("The Frameworks tab presents a left-detail / right-tree layout: select any control group to "
     "expand its items, see the assessed maturity (CMMI) derived from telemetry, and inspect the "
     "crosswalk to the underlying CSF functions. It supports an auditor-style drill and PDF export.")

h2("CFO — Finance")
bullet("total dollar exposure and its concentration.", bold_lead="Exposure:  ")
bullet("modelled tail loss, business-interruption, and the insurance-coverage gap.", bold_lead="Financial posture:  ")
bullet("blended return on security initiatives, derived from posture trend and decision cost/exposure.", bold_lead="ROI:  ")
bullet("SOX ITGC and payment-anomaly signals where connected.", bold_lead="Fraud & controls:  ")

h2("CRO — Chief Risk")
para("Places cyber on one comparable scale beside your other principal risks (credit/market, "
     "operational, compliance, third-party), with a KRI board and mitigation tracking. This is the "
     "view a risk committee uses to weigh cyber against everything else on the enterprise risk register.")

h2("CTO — Technology")
bullet("which systems carry the business, from your live crown-jewel data.", bold_lead="Critical systems:  ")
bullet("application-security posture, critical vulnerabilities, and technical debt.", bold_lead="App-sec health:  ")
bullet("AI inventory and how much of it is governed.", bold_lead="AI systems:  ")

h2("CLO — General Counsel")
bullet("a timed, evidenced determination workbench with a countdown.", bold_lead="Materiality:  ")
bullet("open and overdue DSARs, records of processing, and access to personal data.", bold_lead="Privacy:  ")
bullet("active litigation holds and IP exposure.", bold_lead="Litigation & IP:  ")
bullet("regulatory obligations by jurisdiction, with notification clocks.", bold_lead="Obligations:  ")

h2("COO — Operations")
bullet("business-continuity coverage and process resilience.", bold_lead="Continuity:  ")
bullet("recovery-time and recovery-point objectives (RTO/RPO) versus target.", bold_lead="Recovery:  ")
bullet("backup immutability and identity-recovery readiness.", bold_lead="Backups:  ")

h2("CPO — Product")
bullet("secure-by-design coverage across the product estate.", bold_lead="Secure-by-design:  ")
bullet("open product risks, customer-data handling, and MFA adoption.", bold_lead="Product risk:  ")

h2("Internal Audit — Chief Audit Executive")
para("Presents the audit plan, workpapers, and control assurance mapped to GRC, so audit can trace "
     "any assertion in the platform back to its evidence and coverage.")

h2("Board — Directors")
para("A board-level view of oversight, disclosure readiness, resilience investment, and posture "
     "versus peers — printable as a board/regulator pack (Section 6).")

# ===========================================================================
# 6 · CORE WORKFLOWS
# ===========================================================================
h1("Core Workflows")

h2("Recording and committing a decision")
numbered("1", "Open the decision", "From any seat, select a surfaced decision to see its concrete "
    "pros, cons, cost, and the exposure it addresses.")
numbered("2", "Record your choice", "Enter your rationale and confirm. The decision is stamped with "
    "your name and timestamp.")
numbered("3", "Edit within 24 hours", "For 24 hours the decision remains editable. After that it "
    "commits to the permanent, accountable record.")
numbered("4", "Push to ticketing", "Send a committed decision to Jira or ServiceNow so execution is "
    "tracked in your existing workflow. Status is pulled back on refresh.")

h2("Refreshing live signals")
para("Selecting refresh re-reads your connected tools and updates every live metric, with the "
     "‘as of’ timestamp advancing to match. Illustrative metrics are unaffected. Connected "
     "integrations can also refresh automatically on a schedule set by your administrator.")

h2("Running the peer benchmark")
para("From the CEO or Board seat, the peer benchmark compares your posture to published medians for "
     "your industry. It is drawn from published baselines and requires no data sharing — nothing "
     "about your organization leaves the appliance to produce it.")

h2("Drilling a framework control")
para("In the CISO Frameworks tab, expand a control group, select a control, and inspect its assessed "
     "maturity, the telemetry behind it, and its crosswalk to NIST CSF. Export the assessment as a "
     "PDF for auditors. All mapped control identifiers are validated against the source frameworks.")

h2("Exporting the board pack")
para("The Board seat produces a printable board- and regulator-ready report of oversight, disclosure "
     "readiness, and peer posture. Generate it ahead of a board meeting or a regulatory filing window.")

# ===========================================================================
# 7 · ADMINISTRATION
# ===========================================================================
h1("Administration")

h2("The appliance")
para("Nerion runs as a hardened virtual appliance in your environment. Administrators are "
     "responsible for deploying the image, installing the license file, and setting the small number "
     "of environment controls that govern enforcement. Because the platform is fully offline, routine "
     "operation requires no outbound connectivity.")

h2("Licensing and renewal")
numbered("1", "Read the machine fingerprint", "Each appliance has a unique hardware fingerprint, "
    "visible to the administrator. A license can be bound to it so it cannot be copied to another VM.")
numbered("2", "Install the license", "Place the license file provided by your vendor into the "
    "appliance. The platform validates it cryptographically — it cannot be forged or edited.")
numbered("3", "Watch the countdown", "The platform shows days remaining. A banner prompts renewal "
    "as expiry approaches; a short grace period follows expiry before access is blocked.")
numbered("4", "Renew", "Your vendor issues a fresh license for the new term; installing it extends "
    "the clock with no reinstallation.")

callout("Tamper protection",
    "The appliance protects the vendor’s intellectual property, not your data. If the software "
    "itself is tampered with, it seals and renders its own protected components inert — it never "
    "touches, encrypts, or deletes your business data. An accidental lock is cleared by a vendor "
    "recovery token in minutes.")

h2("Data residency")
para("All customer data remains inside the appliance in your environment for the life of the "
     "subscription. Nerion does not transmit your business data externally to operate, benchmark, "
     "or validate its license.")

# ===========================================================================
# 8 · SUPPORT & FAQ
# ===========================================================================
h1("Support & Frequently Asked Questions")

def faq(q, a):
    qp = doc.add_paragraph(); qp.paragraph_format.space_before = Pt(8); qp.paragraph_format.space_after = Pt(2)
    run(qp, "Q  ", size=11, bold=True, color=CYAN, font=HEAD_FONT); run(qp, q, size=11, bold=True, color=NAVY)
    ap = doc.add_paragraph(); ap.paragraph_format.left_indent = Inches(0.28); ap.paragraph_format.space_after = Pt(4)
    run(ap, a)

faq("A metric shows as “illustrative”. How do I make it live?",
    "Connect the tool that feeds it in onboarding Step 6, or provide the underlying value in the "
    "relevant onboarding section. The metric switches to live on the next refresh.")
faq("Can I change onboarding answers after going live?",
    "Yes. Every input is editable at any time; the cockpit recomputes from your changes. Nothing is locked.")
faq("Where do my decisions go?",
    "They are recorded against your name, committed after 24 hours, and — if you choose — pushed to "
    "Jira or ServiceNow, with status pulled back on refresh.")
faq("Does any of our data leave the environment?",
    "No. The appliance is fully offline for operation, benchmarking, and licensing. Your data stays in your network.")
faq("What happens when the license expires?",
    "A renewal banner appears before expiry; a short grace period follows. Installing a renewed "
    "license restores full access without reinstalling the platform.")
faq("The platform is locked after a change to the VM. What now?",
    "This is the tamper seal. Contact your vendor with the fingerprint and code shown; a signed "
    "recovery token clears it in minutes. Your data is untouched.")

h2("Getting help")
para("For platform questions, contact your Nerion administrator first — most seat and onboarding "
     "questions are answered by reviewing the relevant section above. For licensing, appliance, or "
     "recovery matters, your administrator will engage your Nerion vendor contact.")

# closing rule
end = doc.add_paragraph(); end.paragraph_format.space_before = Pt(18)
_borders(end, ['bottom'], sz=10, color=NAVYHX)
endp = doc.add_paragraph(); endp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(endp, "NERION  ·  Cyber Business Operations Platform", size=9, bold=True, color=BLUE, spacing=30)
endp2 = doc.add_paragraph(); endp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(endp2, "Confidential — Distribution limited to licensed users", size=8, italic=True, color=MUTE)

OUT = "/home/user/Cyber-Rx/Nerion_User_Guide.docx"
doc.save(OUT)
print("saved", OUT)
